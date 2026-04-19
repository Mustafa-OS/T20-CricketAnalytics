"""
Generate CAIS_Summary.docx — a short (1-2 page) reference card for the
Context-Adjusted Impact Score. Intended for readers who want the essentials
without the full mathematical walkthrough in CAIS_Methodology.docx.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours (match the full methodology doc) ───────────────────
AMBER = RGBColor(0xB4, 0x7A, 0x00)
SLATE = RGBColor(0x1F, 0x2A, 0x3A)
MUTED = RGBColor(0x5F, 0x6B, 0x7F)
CODE_BG_HEX = "F3F1EC"

doc = Document()

# ── Page & default styles ──────────────────────────────────────
for section in doc.sections:
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(10.5)

for i in range(1, 3):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = SLATE if i > 1 else AMBER
    h.font.bold = True

# ── Helpers (trimmed from build_cais_doc.py) ──────────────────
def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_formula(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x80, 0x50, 0x00)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), CODE_BG_HEX)
    p_pr.append(shd)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), 'B47A00')
    pBdr.append(left)
    p_pr.append(pBdr)

def add_para(text, bold=False, italic=False, size=10.5, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def add_bullet(text):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade_cell(hdr[i], "B47A00")
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
#   TITLE
# ═══════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CAIS — Quick Reference')
r.font.size = Pt(28); r.font.bold = True
r.font.color.rgb = AMBER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Context-Adjusted Impact Score  ·  v3  ·  one-page summary')
r.font.size = Pt(11); r.font.italic = True
r.font.color.rgb = MUTED

add_para(
    "This is the short version. See CAIS_Methodology.docx for the full "
    "derivation, rationale, and worked examples.",
    italic=True, color=MUTED, size=9.5
)

# ═══════════════════════════════════════════════════════════════
#   1. THE IDEA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1.  The idea', level=1)
add_para(
    "Conventional T20 averages and strike rates flatten context. A 35 off "
    "25 in the 18th over defending 160 is not the same as a 35 off 25 in a "
    "dead middle over against the 5th bowler. CAIS re-prices every legal "
    "delivery by nine context factors — phase, bowler role, batter quality, "
    "form, pressure, partnership broken, early-wicket flag, tournament stage, "
    "and (in World Cups) opponent quality — and rolls them up into per-player "
    "batting and bowling scores."
)

# ═══════════════════════════════════════════════════════════════
#   2. FORMULAS
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2.  The two formulas', level=1)
add_para("Batting — per ball faced:", bold=True)
add_formula(
    "CAIS_bat(b) = [ Σ ( runs · phase_bat · pressure · stage · opp_bat ) / balls ] · 100 · form_avg"
)
add_para("Bowling — per over bowled:", bold=True)
add_formula(
    "wicket_value = 30 · phase×role · batter_tier · form · pressure\n"
    "                  · partnership · early_wicket · stage · opp_bowl\n"
    "run_cost     = runs_conceded · phase_bowl · 0.5 · stage · opp_bowl\n"
    "CAIS_bowl(B) = Σ ( is_wicket · wicket_value − run_cost ) / overs"
)
add_para(
    "Every multiplier above is bounded; several are competition- or "
    "situation-specific and fall back to 1.0 when not applicable (stage "
    "only in tournament knockouts, opponent only at the World Cup).",
    italic=True, color=MUTED, size=10
)

# ═══════════════════════════════════════════════════════════════
#   3. WEIGHT TABLE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3.  The nine factors, at a glance', level=1)
add_table(
    headers=['Factor', 'Range', 'Applies to', 'What it captures'],
    rows=[
        ['phase_bat',         '0.95 – 1.35', 'batting',        'PP easiest, death hardest'],
        ['phase_bowl',        '1.00 – 1.30', 'bowling cost',   'PP and death cost more than middle'],
        ['phase × role',      '1.20 – 2.00', 'bowling wicket', 'pace-PP wickets worth most (2.0×)'],
        ['batter_tier',       '0.75 – 1.50', 'bowling wicket', 'quartile tier from career avg + SR'],
        ['form_mult',         '1.00 – 1.45', 'both',           'rolling 3-inns runs vs league mean'],
        ['pressure',          '1.00 – 1.50', 'both',           'wicket-loss × 2nd-inns RRR, compounded'],
        ['partnership_mult',  '1.00 – 1.60', 'bowling wicket', 'bigger stand broken = bigger bonus'],
        ['early_wicket_mult', '1.00 – 1.35', 'bowling wicket', 'first wicket in first over worth 1.35×'],
        ['stage_mult',        '1.00 – 1.30', 'both',           'group → QF → SF → final ladder'],
        ['opp_mult (WC only)','0.85 – 1.20', 'both',           'Full Member ↔ associate asymmetry'],
    ],
    col_widths=[3.6, 2.4, 2.8, 7.5]
)

# ═══════════════════════════════════════════════════════════════
#   4. PHASE WEIGHTS DETAIL
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4.  Phase weights', level=1)
add_table(
    headers=['Phase', 'Overs', 'phase_bat (runs scored)', 'phase_bowl (runs conceded cost)'],
    rows=[
        ['Powerplay', '1–6',   '0.95', '1.20'],
        ['Middle',    '7–15',  '1.15', '1.00'],
        ['Death',     '16–20', '1.35', '1.30'],
    ],
    col_widths=[3.5, 2.5, 4.5, 5.5]
)
add_para(
    "Inverted on the batting side from most public metrics: the powerplay is "
    "the easiest phase to score in (fielding restrictions, hard new ball, "
    "containment-first plans), so runs there get a sub-unit weight."
)

# ═══════════════════════════════════════════════════════════════
#   5. STAGE + OPPONENT
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5.  Stage and opponent', level=1)
add_para("Stage multiplier — applies in every tournament with a knockout bracket (11 leagues + the T20 World Cup). Bilateral T20Is are excluded.", bold=True)
add_table(
    headers=['Stage', 'stage_mult'],
    rows=[
        ['Group',                           '1.00'],
        ['Quarter / playoff bubble',        '1.10'],
        ['Semi / Qualifier / Eliminator',   '1.20'],
        ['Final',                           '1.30'],
    ],
    col_widths=[9, 3]
)
add_para("Opponent-quality multiplier — World Cups only. Rewards underdogs, softens credit for favourites. Everywhere else = 1.0.", bold=True)
add_table(
    headers=['Matchup', 'bat_opponent', 'bowl_opponent'],
    rows=[
        ['Full Member vs Full Member',          '1.00', '1.00'],
        ['Full Member batter vs associate',     '0.85', '—'],
        ['Associate batter vs Full Member',     '1.20', '—'],
        ['Full Member bowler vs associate',     '—',    '0.85'],
        ['Associate bowler vs Full Member',     '—',    '1.20'],
    ],
    col_widths=[7, 3, 3]
)
add_para(
    "The 12 Full Members: Afghanistan, Australia, Bangladesh, England, "
    "India, Ireland, New Zealand, Pakistan, South Africa, Sri Lanka, "
    "West Indies, Zimbabwe. Everyone else in the WC is an associate.",
    italic=True, color=MUTED, size=10
)

# ═══════════════════════════════════════════════════════════════
#   6. INTERPRETATION CUES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6.  How to read a CAIS number', level=1)
add_bullet('CAIS > raw SR → player performs in hard phases or under pressure; probably elevated by death-overs runs, chase-under-collapse, or a knockout bump.')
add_bullet('CAIS ≈ raw SR → context-neutral player whose raw numbers already reflect their value.')
add_bullet('CAIS < raw SR → scoring concentrated in soft phases, low-pressure middle overs, or frequently out of form.')
add_bullet('For bowling CAIS, a positive value means net wicket value exceeds run cost per over; negative means they\'re leaking more than they take.')

# ═══════════════════════════════════════════════════════════════
#   7. WHAT'S NEW IN V3
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7.  What v3 adds over v2', level=1)
add_bullet('Stage multiplier (group → final ladder), inferred from match dates because Cricsheet does not tag knockouts explicitly.')
add_bullet('Stage restricted to the 11 supported leagues + WC. T20Is are exempt — the last match of a calendar year is not a "final".')
add_bullet('Stage map keyed on (competition, match_id) so WC knockout bumps do not leak into the T20I bilaterals view.')
add_bullet('WC-only opponent-quality multiplier: 0.85 when a Full Member plays an associate, 1.20 the other way.')
add_bullet('Both stage and opponent apply to the bowling run_cost side as well, not just wicket_value — leaking runs in a final is proportionally more costly.')

# ═══════════════════════════════════════════════════════════════
out_path = 'CAIS_Summary.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
