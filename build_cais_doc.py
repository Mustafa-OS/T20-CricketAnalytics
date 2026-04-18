"""
Generate CAIS_Methodology.docx — full mathematical write-up of the
Context-Adjusted Impact Score engine.
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Colours (dark accent tones) ────────────────────────────────
AMBER = RGBColor(0xB4, 0x7A, 0x00)
SLATE = RGBColor(0x1F, 0x2A, 0x3A)
MUTED = RGBColor(0x5F, 0x6B, 0x7F)
CODE_BG_HEX = "F3F1EC"

doc = Document()

# ── Page & default styles ──────────────────────────────────────
for section in doc.sections:
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

# heading styles
for i in range(1, 4):
    h = doc.styles[f'Heading {i}']
    h.font.name = 'Calibri'
    h.font.color.rgb = SLATE if i > 1 else AMBER
    h.font.bold = True

# ── Helper functions ───────────────────────────────────────────
def shade_cell(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def add_formula(text):
    """Insert a monospace boxed formula line."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor(0x80, 0x50, 0x00)

    # Light background shading via paragraph properties
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), CODE_BG_HEX)
    p_pr.append(shd)
    # left border
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'), 'single')
    left.set(qn('w:sz'), '18')
    left.set(qn('w:space'), '6')
    left.set(qn('w:color'), 'B47A00')
    pBdr.append(left)
    p_pr.append(pBdr)

def add_para(text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color: run.font.color.rgb = color
    return p

def add_rich_para(segments):
    """segments: list of (text, {bold, italic, mono, color})."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    for text, opts in segments:
        run = p.add_run(text)
        run.bold   = opts.get('bold', False)
        run.italic = opts.get('italic', False)
        if opts.get('mono'):
            run.font.name = 'Consolas'
            run.font.size = Pt(10.5)
        if opts.get('color'):
            run.font.color.rgb = opts['color']
    return p

def add_bullet(text, rich=None):
    if rich:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        for t, opts in rich:
            r = p.add_run(t)
            r.bold   = opts.get('bold', False)
            r.italic = opts.get('italic', False)
            if opts.get('mono'):
                r.font.name = 'Consolas'
                r.font.size = Pt(10.5)
        return p
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    return p

def add_table(headers, rows, col_widths=None, highlight_header=True):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    if col_widths:
        for i, w in enumerate(col_widths):
            for cell in t.columns[i].cells:
                cell.width = Cm(w)
    # header
    hdr_cells = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ''
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        if highlight_header:
            shade_cell(hdr_cells[i], "B47A00")
    # rows
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = ''
            p = cells[ci].paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(10.5)
    doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
#   TITLE PAGE HEADER
# ═══════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('CAIS')
r.font.size = Pt(42); r.font.bold = True
r.font.color.rgb = AMBER

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Context-Adjusted Impact Score')
r.font.size = Pt(18); r.font.bold = True
r.font.color.rgb = SLATE

sub2 = doc.add_paragraph()
sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Full Mathematical Specification & Logic  ·  v2')
r.font.size = Pt(12); r.font.italic = True
r.font.color.rgb = MUTED

ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ver.add_run(
    'v2 adds chase-pressure (2nd-innings RRR), flipped batting-phase weights '
    '(powerplay easiest, death hardest), partnership-broken multiplier, '
    'and early-wicket bonus.'
)
r.font.size = Pt(10); r.font.italic = True
r.font.color.rgb = MUTED

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════
#   1. OVERVIEW
# ═══════════════════════════════════════════════════════════════
doc.add_heading('1.  Overview', level=1)
add_para(
    "CAIS is a per-delivery weighted metric computed over ball-by-ball T20 data. "
    "Each legal delivery contributes a score that depends not only on its raw outcome "
    "(runs scored, wicket taken, runs conceded) but on up to SEVEN measurable context signals: "
    "the phase of the innings, the bowler's role (pace or spin), the quality of the batter "
    "on strike, the batter's recent form, the in-match pressure state (itself a combination "
    "of wicket-loss pressure and 2nd-innings chase pressure), the size of the partnership "
    "that just got broken (for wickets), and whether the wicket is an early breakthrough. "
    "Batting CAIS and Bowling CAIS are reported separately."
)

add_para(
    "The design philosophy is to reward outcomes that a coach or analyst would intuitively "
    "rate higher — a powerplay wicket against an elite in-form batter, or runs scored during "
    "a middle-order collapse — without relying on any external labels. Every weight is "
    "derived directly from the ball-by-ball dataset itself."
)

# ═══════════════════════════════════════════════════════════════
#   2. NOTATION
# ═══════════════════════════════════════════════════════════════
doc.add_heading('2.  Notation', level=1)
add_para("The CAIS engine operates on the set L of legal deliveries "
         "(wides are excluded because they don't count as balls faced). For a delivery i ∈ L we define:")

add_bullet(None, rich=[('rᵢ', {'mono': True}), (' — runs scored off the bat on delivery i (0, 1, 2, 3, 4 or 6).', {})])
add_bullet(None, rich=[('Rᵢ', {'mono': True}), (' — total runs conceded (bat + extras).', {})])
add_bullet(None, rich=[('wᵢ', {'mono': True}), (' — wicket indicator (1 if wicket fell, else 0).', {})])
add_bullet(None, rich=[('oᵢ', {'mono': True}), (' — 1-indexed over number of delivery i (re-indexed if source data uses 0–19).', {})])
add_bullet(None, rich=[('mᵢ', {'mono': True}), (' — match_id × inning key for delivery i.', {})])
add_bullet(None, rich=[('bᵢ, Bᵢ', {'mono': True}), (' — batter on strike and bowler for delivery i.', {})])

# ═══════════════════════════════════════════════════════════════
#   3. PHASE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('3.  Phase Classification', level=1)
add_para(
    "Every ball is bucketed into one of three innings phases by over number. The phase "
    "produces two separate multipliers: one applied to runs scored by a batter (phase_bat), "
    "and one applied to runs conceded by a bowler (phase_bowl). They are not symmetric — "
    "the batting gradient is now inverted from v1 to reflect real T20 difficulty."
)
add_formula("phase(o) = powerplay   if  o ≤ 6\n"
            "         = death       if  o ≥ 16\n"
            "         = middle      otherwise")
add_table(
    headers=['Phase', 'Overs', 'phase_bat', 'phase_bowl'],
    rows=[
        ['Powerplay', '1–6',  '0.95', '1.20'],
        ['Middle',    '7–15', '1.15', '1.00'],
        ['Death',     '16–20','1.35', '1.30'],
    ],
    col_widths=[4, 3, 4, 4]
)
add_para(
    "Rationale — batting (updated in v2): scoring in the powerplay is typically the easiest "
    "phase of a T20 innings. Only two fielders are allowed outside the 30-yard circle, the "
    "ball is hardest and travels fastest off the bat, and most bowlers still err on the side "
    "of containment with the new ball. A run in the powerplay therefore gets a sub-unit "
    "weight (0.95). Middle overs are harder (spinners, deeper fields, accumulation), so they "
    "get 1.15. Death-overs runs are the hardest — yorkers, slower balls, full spread field, "
    "death specialists bowling to plans — hence the premium at 1.35.", italic=False
)
add_para(
    "Rationale — bowling: the bowling weights are asymmetric from batting. They price the "
    "COST of runs conceded, so a death-over boundary hurts more (1.3), a middle-overs single "
    "is neutral (1.0), and a powerplay boundary costs more than middle because the opposition "
    "is expected to score there (1.2). This keeps the bowler's cost curve aligned with match impact.", italic=False
)

# ═══════════════════════════════════════════════════════════════
#   4. BOWLER ROLE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('4.  Bowler Role Inference', level=1)
add_para(
    "The dataset contains no pace/spin label, so role is inferred from where each bowler "
    "is actually used. Spinners overwhelmingly bowl in the middle overs (7–15); pacers do "
    "the new-ball and death work."
)
add_formula("mid_pct(B) = (# balls bowled by B in overs 7–15) / (total balls bowled by B)\n"
            "role(B)    = 'spin'  if  mid_pct(B) > 0.55\n"
            "           = 'pace'  otherwise")
add_para(
    "The 0.55 threshold was chosen because typical pacers bowl ~40% of their overs in "
    "the middle phase (9 of ~20 possible overs), while specialist spinners land above "
    "60–70%. A bowler who moves freely between pace and spin roles rarely exists at PSL "
    "level, so misclassification is negligible.", italic=False
)

# ═══════════════════════════════════════════════════════════════
#   5. BATTER QUALITY TIER
# ═══════════════════════════════════════════════════════════════
doc.add_heading('5.  Batter Quality Tier', level=1)
add_para(
    "Every batter with at least 5 career innings is assigned a quality multiplier that "
    "captures both how many runs they score per dismissal (average) and how fast they "
    "score (strike rate). Using both avoids rewarding tailenders with inflated averages "
    "from not-outs, and avoids rewarding pinch-hitters with high SR but low bulk runs."
)
doc.add_heading('5.1  Normalisation', level=2)
add_para("For each batter we compute career average avg_b and career strike rate sr_b from legal deliveries, then min-max normalise both across all qualifying batters:")
add_formula("avg_n(b) = (avg_b − min(avg)) / (max(avg) − min(avg) + ε)\n"
            "sr_n(b)  = (sr_b  − min(sr))  / (max(sr)  − min(sr)  + ε)")
add_para("ε = 1e-9 prevents division by zero if every batter somehow shares the same value. "
         "Both normalised values lie in [0, 1].")

doc.add_heading('5.2  Composite Score', level=2)
add_formula("score(b) = 0.6 · avg_n(b) + 0.4 · sr_n(b)")
add_para("Average is weighted more heavily than strike rate because a high average is a "
         "stronger signal of consistent contribution, while SR without volume can be noisy.")

doc.add_heading('5.3  Tier Mapping', level=2)
add_para("The composite score is then converted into a discrete multiplier using its own "
         "quartile cutoffs (p25, p50, p75 of score across all qualifying batters):")
add_table(
    headers=['Percentile', 'Tier', 'Multiplier'],
    rows=[
        ['≥ 75th', 'Elite',       '1.50'],
        ['50–75',  'Good',        '1.20'],
        ['25–50',  'Average',     '1.00'],
        ['< 25th', 'Lower-tier',  '0.75'],
    ],
    col_widths=[4, 4, 4]
)
add_para("Effect: dismissing a top-quartile batter is worth twice as much as dismissing "
         "a bottom-quartile one (1.50 ÷ 0.75 = 2.0). The tiers are also recomputed whenever "
         "the dataset is updated, so they're self-calibrating as the league evolves.")

# ═══════════════════════════════════════════════════════════════
#   6. FORM MULTIPLIER
# ═══════════════════════════════════════════════════════════════
doc.add_heading('6.  Form Multiplier (Detailed)', level=1)
add_para(
    "The form multiplier captures how in-form a batter was when a given delivery was "
    "bowled. It sits in the range [1.00, 1.45] and is applied two ways: for batting CAIS "
    "it scales the batter's own career contribution (as a career-average of form values); "
    "for bowling CAIS it is attached to each delivery based on the striker on that ball."
)

doc.add_heading('6.1  Step 1 — Per-innings run totals', level=2)
add_para("Legal deliveries are aggregated to one row per (batter, match_id, date) with a single column: runs scored in that innings.")
add_formula("innings(b, m) = Σ  rᵢ    for every ball i in match m faced by batter b\n"
            "                i∈L_bm")

doc.add_heading('6.2  Step 2 — Rolling 3-innings mean (shifted)', level=2)
add_para("For each batter, innings are sorted chronologically. A rolling mean of the runs column is computed over the previous 3 innings (not including the current one):")
add_formula("form_runs(b, m) = mean( innings(b, m−1), innings(b, m−2), innings(b, m−3) )")
add_para("Implementation uses pandas:", italic=False)
add_formula(".groupby('batter')['batsman_runs']\n"
            ".transform(lambda x: x.shift(1).rolling(3, min_periods=1).mean())")
add_rich_para([
    ('The ', {}), ('shift(1)', {'mono': True}),
    (' call is critical: it excludes the current innings, so form is strictly historical. ', {}),
    ('min_periods=1', {'mono': True}),
    (' means the rolling average starts producing values from the second innings onwards — before then, the value is NaN and the default multiplier of 1.0 is used.', {}),
])

doc.add_heading('6.3  Step 3 — Mapping to a [1.00, 1.45] multiplier', level=2)
add_para("Let μ_form be the global mean of form_runs across every batter-innings row in the dataset. For each delivery (or innings), the raw form ratio is:")
add_formula("ratio = form_runs(b, m) / μ_form")
add_para("The ratio is then linearly shifted and clipped:")
add_formula("form_mult = clip( 0.85 + 0.6 · ratio,  lower = 1.00,  upper = 1.45 )")

add_para("Concrete values:", bold=True)
add_table(
    headers=['Ratio (form vs league avg)', 'Raw 0.85 + 0.6·ratio', 'After clip → form_mult'],
    rows=[
        ['0.0  (3 consecutive ducks)',   '0.85', '1.00'],
        ['0.25',                          '1.00', '1.00'],
        ['0.5',                           '1.15', '1.15'],
        ['1.0 (average)',                 '1.45', '1.45'],
        ['1.5',                           '1.75', '1.45'],
        ['2.0 (twice the league avg)',    '2.05', '1.45'],
    ],
    col_widths=[5.5, 4.0, 4.5]
)

add_para("Why this shape?", bold=True)
add_bullet('Floor at 1.00 means a cold-streak batter still counts for at least their baseline value — the metric never punishes you for being out of form, it only rewards being in form.')
add_bullet('Ceiling at 1.45 prevents a single purple-patch innings from distorting the leaderboard. A batter on a 3-innings 60-80-90 streak is not three times as valuable as a 30-40-50 one.')
add_bullet('Linear slope of 0.6 means each 10% improvement over league average yields a 0.06 gain, so the multiplier climbs smoothly from 1.00 at ratio=0.25 up to 1.45 at ratio=1.0. Gradual, not binary.')
add_bullet('Offset of 0.85 sets the zero-form floor so that a batter with exactly 25% of the league\'s average form gets a flat 1.0 (i.e., no bonus, no penalty). Anything below that is floored.')

doc.add_heading('6.4  Application in Batting CAIS', level=2)
add_para(
    "Because batting CAIS is computed by aggregating every ball of a player's career (not innings-by-innings), we collapse the per-match form values into a single career multiplier:"
)
add_formula("form_avg(b) = mean{ form_mult(b, m) : m ∈ matches played by b }")
add_para("This is then multiplied in exactly once at the end of batting CAIS.")

doc.add_heading('6.5  Application in Bowling CAIS', level=2)
add_para("Per-ball — the bowler who dismisses an in-form Babar Azam on that particular day gets credit; dismissing the same Babar during a cold streak pays less. Implementation: each delivery i looks up form_mult(bᵢ, mᵢ).")

# ═══════════════════════════════════════════════════════════════
#   7. PRESSURE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('7.  Pressure Index (Detailed)', level=1)
add_para(
    "Pressure captures the in-match state at the exact moment of the delivery. v2 "
    "decomposes pressure into TWO independent components that are then combined: "
    "(A) wicket-loss pressure — the collapse-in-progress signal, applicable in both "
    "innings; and (B) chase pressure — the required-run-rate signal in the second "
    "innings only. Each is in [1.00, 1.30]; they multiply together and are then "
    "clipped to a final range of [1.00, 1.50]."
)
add_formula("pressure(i) = clip( wicket_pressure(i) · chase_pressure(i),  1.00,  1.50 )")

# ── 7A. Wicket-loss pressure
doc.add_heading('7A.  Wicket-loss Pressure', level=2)
add_para("Deliveries are ordered by (match_id, inning, over, ball). Within each match-inning, cumulative wickets are tracked:")
add_formula("cum_wkts(i) = Σ  w_j    for every delivery j in the same match-inning at or before i\n"
            "              j≤i")
add_para("A raw signal combines wickets already lost with how late in the innings we are:")
add_formula("raw(i) = clip( cum_wkts(i) / max(oᵢ, 1) · 2,  lower = 0,  upper = 3 )")
add_para("Dividing by the over number softens early wickets. Hard-clipping at 3 prevents pathological states (e.g. 7/3 in the 2nd over) from dominating. The raw signal is then mapped:")
add_formula("wicket_pressure(i) = clip( 1 + 0.1 · raw(i),  lower = 1.00,  upper = 1.30 )")

add_para("Worked examples:", bold=True)
add_table(
    headers=['Match state', 'cum_wkts', 'over', 'raw', 'wicket_pressure'],
    rows=[
        ['Fresh start, ball 1',            '0', '1',  '0.00', '1.00'],
        ['1 down, over 5',                 '1', '5',  '0.40', '1.04'],
        ['3 down, over 10 (middle wobble)','3', '10', '0.60', '1.06'],
        ['6 down, over 14 (collapse)',     '6', '14', '0.86', '1.09'],
        ['7 down, over 11 (heavy collapse)','7','11', '1.27', '1.13'],
        ['8 down, over 8 (extreme)',       '8', '8',  '2.00', '1.20'],
        ['9 down, over 6 (pathological)',  '9', '6',  '3.00', '1.30'],
    ],
    col_widths=[6.5, 2.5, 2.0, 2.0, 3.2]
)

# ── 7B. Chase pressure (NEW)
doc.add_heading('7B.  Chase Pressure (2nd-innings only)', level=2)
add_para(
    "Batting second in T20 is a fundamentally different cognitive task: you know the exact target, the scoreboard counts down, and a rising required run rate compresses decision-making. v2 quantifies this via the required run rate at the instant each delivery is bowled."
)
add_para("Step 1 — First-innings target.", bold=True)
add_formula("target(m) = Σ  total_runs(i)    for every ball i in inning 1 of match m    +  1\n"
            "            i∈L¹_m")
add_para("Step 2 — Running state at delivery i in inning 2.", bold=True)
add_formula("runs_scored_so_far(i) = cumsum of total_runs in inning 2 of match m up to — but not including — ball i\n"
            "runs_needed(i)        = max( target(m) − runs_scored_so_far(i),  0 )\n"
            "balls_bowled(i)       = (oᵢ − 1) · 6 + ballᵢ\n"
            "balls_remaining(i)    = max( 120 − (balls_bowled(i) − 1),  1 )\n"
            "required_rr(i)        = runs_needed(i) / (balls_remaining(i) / 6)")
add_para("Step 3 — Map required RR to multiplier.", bold=True)
add_formula("chase_pressure(i) = clip( 1 + (required_rr(i) − 8.0) · 0.04,  lower = 1.00,  upper = 1.30 )")

add_para("Concrete values:", bold=True)
add_table(
    headers=['Situation', 'required_rr', 'chase_pressure'],
    rows=[
        ['Coasting home (easy chase)',     '5.0',  '1.00'],
        ['Comfortable',                     '8.0',  '1.00'],
        ['Slightly tough',                  '10.0', '1.08'],
        ['Hard',                            '12.0', '1.16'],
        ['Very hard',                       '14.0', '1.24'],
        ['Near-impossible',                 '16.0', '1.30'],
    ],
    col_widths=[7, 4, 4]
)

add_para("Design choices for chase_pressure:", bold=True)
add_bullet('Neutral point at RRR = 8.0 — this is roughly the "par" T20 scoring rate in PSL, above which the chase starts to feel tight.')
add_bullet('Linear slope of 0.04 per extra run of RRR: each +2.5 RRR adds roughly +0.1 to the multiplier.')
add_bullet('Cap at 1.30 and floor at 1.00 keep the chase factor bounded and symmetric with the wicket-loss factor.')
add_bullet('For 1st innings and match-context rows where chase is not applicable, chase_pressure = 1.0 (neutral).')
add_bullet('Balls-remaining clipped to ≥ 1 to avoid divide-by-zero at innings end.')

# ── 7C combined
doc.add_heading('7C.  Combined Pressure', level=2)
add_formula("pressure(i) = clip( wicket_pressure(i) · chase_pressure(i),  1.00,  1.50 )")
add_para("Multiplying compounds both effects — a chase in disarray (lots of wickets down AND a climbing RRR) produces the highest values. The 1.50 cap limits total pressure to at most a 50% context uplift per ball, which is tight enough to stay believable but wide enough to meaningfully separate heroic moments.")

add_para("Combined examples:", bold=True)
add_table(
    headers=['Situation', 'wicket_pressure', 'chase_pressure', 'combined'],
    rows=[
        ['1st innings, settled batting', '1.00', '1.00', '1.00'],
        ['Middle wobble, 1st inns',      '1.09', '1.00', '1.09'],
        ['Comfortable chase',            '1.04', '1.00', '1.04'],
        ['Hard chase, wickets intact',   '1.04', '1.16', '1.21'],
        ['Hard chase + collapse',        '1.15', '1.20', '1.38'],
        ['Extreme (all factors high)',   '1.25', '1.28', '1.50'],
    ],
    col_widths=[6, 3.5, 3.5, 3]
)

doc.add_heading('7D.  Applied in both CAIS flavours', level=2)
add_para("For batting, pressure multiplies runs: a boundary during a hard chase-collapse is worth more. For bowling, pressure multiplies wicket value, but does NOT scale run cost — bleeding runs under pressure isn't doubly punished.")

# ═══════════════════════════════════════════════════════════════
#   8. PHASE × ROLE MATRIX
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8.  Phase × Role Wicket Multiplier', level=1)
add_para(
    "This 2×3 matrix is the single most important hand-calibrated component of bowling "
    "CAIS. It assigns a multiplier to the base wicket value depending on both the bowler's "
    "role and the phase in which the wicket fell. It encodes domain knowledge that coaches "
    "agree on: pace-powerplay wickets are the most valuable single event in T20, and spinners "
    "earn their keep in the middle overs."
)
add_table(
    headers=['Phase', 'Pace', 'Spin'],
    rows=[
        ['Powerplay (1–6)', '2.0×', '1.5×'],
        ['Middle (7–15)',   '1.2×', '1.5×'],
        ['Death (16–20)',   '1.8×', '1.2×'],
    ],
    col_widths=[5.5, 4.5, 4.5]
)
add_bullet(None, rich=[('Pace / Powerplay = 2.0×.', {'bold': True}),
                       (' Highest weight. Pace powerplay wickets are both rare and momentum-flipping. Dismissing a top-order batter in the first 6 overs with the new ball swinging forces a new batter into an attacking field and often collapses a scoring plan.', {})])
add_bullet(None, rich=[('Pace / Death = 1.8×.', {'bold': True}),
                       (' Second highest. Wickets at the death break intent-to-score, but batters are already slogging so the wicket is somewhat more common.', {})])
add_bullet(None, rich=[('Pace / Middle = 1.2×.', {'bold': True}),
                       (' Not the pacer\'s natural home; wickets here are somewhat incidental.', {})])
add_bullet(None, rich=[('Spin / Middle = 1.5×.', {'bold': True}),
                       (' Spinners are supposed to take wickets in the middle — set batters, rotating strike, low risk tolerance. A middle-overs spin breakthrough is exactly the job.', {})])
add_bullet(None, rich=[('Spin / Powerplay = 1.5×.', {'bold': True}),
                       (' Brave captaincy move. Spin in the powerplay against a new ball is high-risk, so wickets there are doubly creditable.', {})])
add_bullet(None, rich=[('Spin / Death = 1.2×.', {'bold': True}),
                       (' Least weight — spinners at the death are often being milked, so wickets are partly the batters giving it away.', {})])

# ═══════════════════════════════════════════════════════════════
#   8.5  PARTNERSHIP-BROKEN MULTIPLIER  (v2)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8.5  Partnership-Broken Multiplier', level=1)
add_para(
    "v2 adds a multiplicative bonus on every WICKET ball according to the size of the "
    "partnership that was just broken. A bowler who dismisses a batter after a 70-run stand "
    "has just erased more match equity than one who nicks off a new batter immediately."
)

doc.add_heading('8.5.1  Tracking partnership size', level=2)
add_para("Within each (match_id, inning) group, we compute the id of the currently-live partnership at the moment BEFORE each ball:")
add_formula("partnership_key(i) = (cumulative wickets in innings up to and including balls j<i)\n"
            "                   = cumsum(w) .shift(1) .fillna(0)    within (match_id, inning)")
add_para("On a wicket ball itself, this key equals the id of the partnership that will be broken by that same ball. Grouping by this key and cumulative-summing total_runs (then subtracting the current ball's runs) yields the size of the partnership as it stood immediately before the delivery:")
add_formula("partnership_runs(i) = Σ  total_runs(j)    minus    total_runs(i)\n"
            "                     j in same partnership_key, j ≤ i")

doc.add_heading('8.5.2  Mapping to a multiplier', level=2)
add_formula("partnership_mult(i) = clip( 1 + (partnership_runs(i) − 20) · 0.012,  1.00,  1.60 )")
add_table(
    headers=['Partnership size (runs)', 'Raw 1 + (x − 20)·0.012', 'After clip'],
    rows=[
        ['10 (just started)',      '0.88', '1.00'],
        ['20 (neutral point)',     '1.00', '1.00'],
        ['30',                     '1.12', '1.12'],
        ['50 (meaningful)',        '1.36', '1.36'],
        ['70 (big stand)',         '1.60', '1.60'],
        ['100 (massive)',          '1.96', '1.60'],
    ],
    col_widths=[5.5, 4.5, 4]
)
add_para("Design notes:", bold=True)
add_bullet('Neutral point at 20 runs — below this, the partnership is still being built and the wicket isn\'t particularly "crucial". The multiplier stays at 1.0.')
add_bullet('Linear slope of 0.012 per extra run — each 10 runs in the partnership adds 0.12 to the multiplier. A 50-run stand → 1.36×.')
add_bullet('Ceiling of 1.60 keeps the bonus bounded even for century partnerships, to stop one freak event from swinging the leaderboard.')
add_bullet('The multiplier is applied to wicket_value, so it has no effect on balls where no wicket fell.')

# ═══════════════════════════════════════════════════════════════
#   8.6  EARLY-WICKET MULTIPLIER  (v2)
# ═══════════════════════════════════════════════════════════════
doc.add_heading('8.6  Early-Wicket Multiplier', level=1)
add_para(
    "A bowler who dismisses an opener in the first over of a T20 innings has created "
    "outsized value — they expose a new batter to a fresh ball immediately, disrupt the "
    "opposition's batting plan, and shift win probability by several points. v2 encodes "
    "this with a small dedicated multiplier."
)
doc.add_heading('8.6.1  Definition', level=2)
add_para("This multiplier applies ONLY to the FIRST wicket of an innings (identified via partnership_key == 0 on the wicket ball itself), and only if it falls early:")
add_table(
    headers=['Condition', 'balls_in_innings', 'early_wicket_mult'],
    rows=[
        ['First wicket in the first over',    '1–6',    '1.35'],
        ['First wicket in overs 2–3',         '7–18',   '1.15'],
        ['Any other wicket',                  'n/a',    '1.00'],
    ],
    col_widths=[6.5, 3.5, 3.5]
)
add_para("balls_in_innings is the 1-indexed ball counter within the innings, computed as (over − 1) · 6 + ball.")

doc.add_heading('8.6.2  Interaction with other factors', level=2)
add_para(
    "The early-wicket multiplier stacks with partnership_mult, phase×role, batter tier, "
    "form, and pressure. Because the partnership size at the moment of a very early wicket "
    "is tiny (<10 runs), partnership_mult will itself be 1.0 on such balls — so the early "
    "multiplier is what meaningfully rewards them. Conversely, a dismissal that is both the "
    "first wicket AND breaks a long opening stand (e.g. 50 runs in 5 overs) lights up "
    "multiple bonuses simultaneously, which is correct behaviour."
)

# ═══════════════════════════════════════════════════════════════
#   9. BATTING CAIS FORMULA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('9.  Batting CAIS — Full Formula', level=1)
add_para("For a batter b with the set of legal deliveries L_b they have faced:")
add_formula("             Σ  ( rᵢ · phase_bat(oᵢ) · pressure(i) )\n"
            "            i∈L_b\n"
            "CAIS_bat(b) = ─────────────────────────────────────── · 100 · form_avg(b)\n"
            "                          |L_b|")

add_para("In words: for each ball the batter has faced, multiply the runs scored by the phase-bat weight and the pressure index; sum over the whole career (or season); divide by total balls faced to put it on a per-ball basis; multiply by 100 to express it on a strike-rate scale; finally apply the batter's career-averaged form multiplier once at the end.")

add_para("Interpretation cues:", bold=True)
add_bullet('CAIS > raw SR ⇒ the batter performs disproportionately in hard phases or under pressure, or is chronically in form.')
add_bullet('CAIS < raw SR ⇒ their scoring is concentrated in soft phases with low pressure, and/or they\'re often out of form.')
add_bullet('CAIS ≈ raw SR ⇒ context-neutral player whose raw numbers already reflect their value.')

add_para("Qualification: min_balls = 50 by default (20 for season-specific leaderboards). This suppresses noise from one-off performances.")

# ═══════════════════════════════════════════════════════════════
#   10. BOWLING CAIS FORMULA
# ═══════════════════════════════════════════════════════════════
doc.add_heading('10.  Bowling CAIS — Full Formula', level=1)
add_para("For a bowler B with legal deliveries L_B they have bowled, we first define a per-ball score.")

doc.add_heading('10.1  Per-ball wicket value (v2 — seven factors)', level=2)
add_formula("wicket_value(i) = 30 · phase_role[role(B), phase(oᵢ)]\n"
            "                    · batter_tier(bᵢ)\n"
            "                    · form_mult(bᵢ, mᵢ)\n"
            "                    · pressure(i)             ← now combined wicket×chase\n"
            "                    · partnership_mult(i)      ← new in v2\n"
            "                    · early_wicket_mult(i)     ← new in v2")
add_para("This value is 0 unless wᵢ = 1 (a wicket fell). The base of 30 calibrates one wicket to roughly one-quarter of a top-phase, elite-batter innings. The seven multiplicative factors each sit in a bounded range, so the final wicket_value is bounded too:")
add_table(
    headers=['Factor', 'Range'],
    rows=[
        ['phase_role',         '1.20 – 2.00'],
        ['batter_tier',        '0.75 – 1.50'],
        ['form_mult',          '1.00 – 1.45'],
        ['pressure (combined)','1.00 – 1.50'],
        ['partnership_mult',   '1.00 – 1.60'],
        ['early_wicket_mult',  '1.00 – 1.35'],
    ],
    col_widths=[6, 4]
)
add_para("Theoretical maximum wicket_value ≈ 30 · 2.00 · 1.50 · 1.45 · 1.50 · 1.60 · 1.35 ≈ 423 — a unicorn ball (pace, powerplay, elite in-form batter dismissed during a big chase collapse after a long partnership, as the first wicket of the innings). Theoretical minimum for a wicket ball ≈ 30 · 1.20 · 0.75 · 1.00 · 1.00 · 1.00 · 1.00 = 27.")

doc.add_heading('10.2  Per-ball run cost', level=2)
add_formula("run_cost(i) = Rᵢ · phase_bowl(oᵢ) · 0.5")
add_para("Every ball incurs a cost proportional to runs conceded (bat + extras), scaled by phase-bowl weight and halved. The 0.5 factor prevents raw run bleeding from dominating the metric relative to the wicket-reward side.")

doc.add_heading('10.3  Aggregation to CAIS', level=2)
add_formula("             Σ  ( wᵢ · wicket_value(i) − run_cost(i) )\n"
            "            i∈L_B\n"
            "CAIS_bowl(B) = ───────────────────────────────────────\n"
            "                        overs(B)")
add_para("where overs(B) = |L_B| / 6. The result is a CAIS-per-over figure — higher is better. Positive CAIS means on average the bowler delivers more wicket value than run cost per over; negative means they're leaking more than they take.")

add_para("Qualification: min_balls = 30 by default (12 for season-specific leaderboards, i.e. two overs minimum).")

# ═══════════════════════════════════════════════════════════════
#   11. WORKED EXAMPLE
# ═══════════════════════════════════════════════════════════════
doc.add_heading('11.  Worked Examples — single wicket balls', level=1)

doc.add_heading('11.1  Huge-impact wicket', level=2)
add_para("Scenario: a pace bowler dismisses an elite in-form opener (partnership so far: 52 runs) in the 3rd over of a 2nd-innings chase where the RRR has already climbed to 12.0 because of a scratchy start.")
add_table(
    headers=['Factor', 'Derivation', 'Value'],
    rows=[
        ['Base',                    'constant',                                              '30'],
        ['phase × role',            'pace × powerplay',                                      '2.00'],
        ['batter_tier',             'top quartile',                                          '1.50'],
        ['form_mult',               'rolling avg 70 vs league mean 35, ratio 2.0 → cap',    '1.45'],
        ['wicket_pressure',         '1 wkt / 3 over · 2 = 0.67 → 1 + 0.07',                  '1.07'],
        ['chase_pressure',          'RRR 12 → 1 + (12−8)·0.04',                              '1.16'],
        ['pressure (combined)',     '1.07 · 1.16',                                           '1.24'],
        ['partnership_mult',        '52 runs → 1 + (52−20)·0.012',                           '1.38'],
        ['early_wicket_mult',       'first wicket, balls_in_innings ≤ 18',                   '1.15'],
    ],
    col_widths=[5, 7, 2.5]
)
add_formula("wicket_value = 30 · 2.00 · 1.50 · 1.45 · 1.24 · 1.38 · 1.15 ≈ 256.8")
add_para("This single ball contributes ~257 to the bowler's running total — more than 8× the base wicket value. That's intentional: it's a rare, league-deciding moment, and CAIS is supposed to surface exactly those events.")

doc.add_heading('11.2  Low-impact wicket', level=2)
add_para("Same raw outcome — a wicket — but the context is benign: a spinner dismissing a tailender in the 18th over of the 1st innings, after a tiny 8-run last-wicket stand.")
add_table(
    headers=['Factor', 'Derivation', 'Value'],
    rows=[
        ['Base',               'constant',                              '30'],
        ['phase × role',       'spin × death',                          '1.20'],
        ['batter_tier',        'below p25',                             '0.75'],
        ['form_mult',          'no history, fallback',                  '1.00'],
        ['wicket_pressure',    '9 wkts / 18 over · 2 = 1.0 → 1+0.10',   '1.10'],
        ['chase_pressure',     '1st innings',                           '1.00'],
        ['pressure (combined)','1.10 · 1.00',                           '1.10'],
        ['partnership_mult',   '8 runs → below 20',                     '1.00'],
        ['early_wicket_mult',  'not first wicket',                      '1.00'],
    ],
    col_widths=[5, 7, 2.5]
)
add_formula("wicket_value = 30 · 1.20 · 0.75 · 1.00 · 1.10 · 1.00 · 1.00 ≈ 29.7")
add_para("A near 9× spread between 11.1 and 11.2 for the same raw event (a dismissal) — which is precisely what 'context-adjusted' is supposed to deliver.")

doc.add_heading('11.3  Run cost on the same balls', level=2)
add_para("Both balls above carried 0 runs (the wicket ball had no runs conceded), so the run-cost side of the ledger was 0. On non-wicket balls, run_cost dominates:")
add_formula("run_cost = runs_conceded · phase_bowl_weight · 0.5\n"
            "  e.g. a powerplay four: 4 · 1.20 · 0.5 = 2.40\n"
            "       a middle-overs single: 1 · 1.00 · 0.5 = 0.50\n"
            "       a death-overs six:     6 · 1.30 · 0.5 = 3.90")

# ═══════════════════════════════════════════════════════════════
#   12. ENGINEERING NOTES
# ═══════════════════════════════════════════════════════════════
doc.add_heading('12.  Engineering & Edge Cases', level=1)
add_bullet('Wides are excluded from L; no-balls are included (the batter can still score).')
add_bullet('Overs are re-indexed to start at 1 if the source data uses 0-indexed overs.')
add_bullet('Pressure uses max(oᵢ, 1) in the denominator to avoid division by zero on the very first ball.')
add_bullet('Form lookups default to 1.0 for any (batter, match_id) missing from the map (e.g. debut innings where no history exists).')
add_bullet('Batter tier lookups default to 1.0 for batters with fewer than 5 career innings, so wickets of unknown batters neither reward nor punish the bowler.')
add_bullet('The _enriched DataFrame is cached on first call; subsequent CAIS computations across seasons reuse it.')

# ═══════════════════════════════════════════════════════════════
#   13. CURRENT CALIBRATION & FUTURE WORK
# ═══════════════════════════════════════════════════════════════
doc.add_heading('13.  Calibration Choices & Future Work', level=1)
add_para(
    "All fixed constants (base = 30, run_cost factor = 0.5, phase weights, phase×role "
    "matrix, form floor/ceiling 1.00–1.45, wicket_pressure floor/ceiling 1.00–1.30, "
    "chase_pressure floor/ceiling 1.00–1.30, combined-pressure cap 1.50, partnership "
    "neutral point 20 with slope 0.012 and cap 1.60, early-wicket multipliers 1.35 / "
    "1.15 / 1.00, tier cutoffs 0.75 / 1.00 / 1.20 / 1.50) were chosen by hand based on "
    "T20 domain intuition and sanity-checked against PSL data. They are deliberately "
    "simple and transparent rather than fitted to an external target."
)
add_para("Changes from v1 → v2:", bold=True)
add_bullet('Batting phase weights flipped: powerplay is now treated as the easiest phase (0.95) and death as the hardest (1.35), replacing v1\'s symmetric 1.2 / 1.0 / 1.3 pattern.')
add_bullet('Pressure decomposed into wicket_pressure × chase_pressure, with chase_pressure (new) driven by 2nd-innings required run rate.')
add_bullet('Combined pressure cap raised from 1.30 to 1.50 to accommodate the compounded signal.')
add_bullet('Partnership-broken multiplier added to wicket value.')
add_bullet('Early-wicket multiplier added to reward first-over breakthroughs.')

add_para("Natural future extensions:", bold=True)
add_bullet('Venue adjustment — the same ball at Rawalpindi (flat) is less impressive than at Karachi (slow).')
add_bullet('Head-to-head calibration — if a specific batter historically dominates a bowler, uplift a wicket from that matchup.')
add_bullet('First-innings "par score" pressure — estimate par from venue history, penalise batting well under par and reward bowling holding a side under it.')
add_bullet('ML-backed tier boundaries — replace quartile tiers with a learned rating (Elo, Glicko) for batter quality.')

add_para("The goal remains explainability: every term in CAIS is a product of transparent, human-interpretable factors. Any ML additions should preserve that property — no black boxes.")

# ═══════════════════════════════════════════════════════════════
out_path = 'CAIS_Methodology.docx'
doc.save(out_path)
print(f'Saved: {out_path}')
